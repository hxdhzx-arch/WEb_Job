"""
app.py — Flask 主入口（工厂模式）
保留所有现有路由，新增 SaaS 后端 Blueprint
"""
import time
import threading
import traceback
from collections import defaultdict
from flask import Flask, render_template, request, jsonify, send_file
from config import PORT, DEBUG, RATE_LIMIT_PER_MINUTE
from services.resume_analyzer import analyze_resume
from services.jd_matcher import match_jd
from services.gemini_client import call_gemini
import io
import docx
from docx.shared import Pt, Inches


def create_app():
    """Flask 应用工厂"""
    app = Flask(__name__)

    # ── 加载配置 ──
    import config as cfg
    app.config["SECRET_KEY"] = cfg.JWT_SECRET_KEY
    app.config["SQLALCHEMY_DATABASE_URI"] = cfg.SQLALCHEMY_DATABASE_URI
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = cfg.SQLALCHEMY_TRACK_MODIFICATIONS
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = cfg.SQLALCHEMY_ENGINE_OPTIONS
    app.config["JWT_SECRET_KEY"] = cfg.JWT_SECRET_KEY
    app.config["JWT_ACCESS_TOKEN_EXPIRES"] = cfg.JWT_ACCESS_TOKEN_EXPIRES
    app.config["JWT_REFRESH_TOKEN_EXPIRES"] = cfg.JWT_REFRESH_TOKEN_EXPIRES

    # ── 初始化扩展 ──
    from backend.extensions import db, jwt, limiter, cors, migrate

    db.init_app(app)
    jwt.init_app(app)
    limiter.init_app(app)
    cors.init_app(app, resources={r"/api/*": {"origins": "*"}})
    migrate.init_app(app, db)

    # ── JWT 错误处理 ──
    @jwt.expired_token_loader
    def expired_token(jwt_header, jwt_payload):
        return jsonify({"success": False, "message": "登录已过期，请重新登录"}), 401

    @jwt.invalid_token_loader
    def invalid_token(error_msg):
        return jsonify({"success": False, "message": "无效的认证令牌"}), 401

    @jwt.unauthorized_loader
    def missing_token(error_msg):
        return jsonify({"success": False, "message": "请先登录"}), 401

    # ── 注册 SaaS API Blueprint ──
    from backend.api import register_blueprints
    register_blueprints(app)

    # ── 确保 data 目录存在 ──
    import os
    os.makedirs(os.path.join(os.path.dirname(__file__), "data"), exist_ok=True)

    # ── 在 app context 中创建表 ──
    with app.app_context():
        # 导入所有模型以注册
        import backend.models  # noqa
        db.create_all()

    # ── 全局错误处理（写入 ErrorLog）──
    @app.errorhandler(500)
    def handle_500(e):
        try:
            from backend.models.error_log import log_error
            log_error(
                message=str(e),
                traceback_str=traceback.format_exc(),
                endpoint=request.path if request else None,
                method=request.method if request else None,
                ip=request.remote_addr if request else None,
            )
        except Exception:
            pass
        return jsonify({"success": False, "message": "服务器内部错误"}), 500

    @app.errorhandler(404)
    def handle_404(e):
        if request.path.startswith("/api/"):
            return jsonify({"success": False, "message": "接口不存在"}), 404
        return render_template("index.html"), 404

    @app.errorhandler(429)
    def handle_429(e):
        return jsonify({"success": False, "message": "请求过于频繁，请稍后再试"}), 429

    # ── 定时任务 ──
    _setup_scheduler(app)

    # ══════════════════════════════════════════
    # 以下是现有路由（完整保留）
    # ══════════════════════════════════════════

    # ── 限流（带自动清理）──
    _request_log = defaultdict(list)
    _log_lock = threading.Lock()

    def _cleanup_stale_ips():
        while True:
            time.sleep(120)
            now = time.time()
            with _log_lock:
                stale = [ip for ip, ts in _request_log.items() if not ts or now - ts[-1] > 120]
                for ip in stale:
                    del _request_log[ip]

    _cleaner = threading.Thread(target=_cleanup_stale_ips, daemon=True)
    _cleaner.start()

    def _is_rate_limited(ip):
        now = time.time()
        with _log_lock:
            _request_log[ip] = [t for t in _request_log[ip] if now - t < 60]
            if len(_request_log[ip]) >= RATE_LIMIT_PER_MINUTE:
                return True
            _request_log[ip].append(now)
        return False

    # ── 页面路由 ──
    @app.route("/")
    def index():
        return render_template("index.html")

    @app.route("/resume")
    def resume_page():
        return render_template("resume.html")

    @app.route("/jd")
    def jd_page():
        return render_template("jd_match.html")

    @app.route("/builder")
    def builder_page():
        return render_template("resume_builder.html")

    @app.route("/career-tools")
    def career_tools_page():
        return render_template("career_tools.html")

    @app.route("/admin-panel")
    def admin_panel():
        return render_template("admin_dashboard.html")

    @app.route("/web-resume")
    def web_resume_page():
        return render_template("web_resume.html")

    @app.route("/personal-site")
    def personal_site_page():
        return render_template("personal_site.html")

    # ── AI 个人网站公开路由 ──
    @app.route("/site/<slug>")
    def personal_site_public(slug):
        from backend.models.personal_site import PersonalSite
        from backend.utils.personal_site_renderer import render_personal_site as _render_site
        site = PersonalSite.query.filter_by(slug=slug, is_published=True).first()
        if not site:
            return render_template("404.html") if _template_exists("404.html") else ("网站不存在", 404)

        # 统计浏览量
        site.view_count = (site.view_count or 0) + 1
        try:
            from backend.extensions import db as _db
            _db.session.commit()
        except Exception:
            pass

        # 服务端渲染
        html = _render_site(site.site_data, site.site_config)
        return html


    # ── 公开预览 API（无需鉴权）──
    @app.route("/api/web-resume/preview-public", methods=["POST"])
    def web_resume_preview_public():
        from backend.utils.web_resume_renderer import render_resume as _render
        data = request.get_json()
        if not data or not data.get("resume_data"):
            return jsonify({"success": False, "message": "缺少数据"}), 400
        try:
            html = _render(data["resume_data"], data.get("web_config"), data.get("template"))
            return jsonify({"success": True, "data": {"html": html}})
        except Exception as e:
            return jsonify({"success": False, "message": str(e)}), 500

    # ── 网页简历公开路由 ──
    @app.route("/r/<slug>")
    def web_resume_public(slug):
        from backend.models.resume import Resume
        from backend.utils.web_resume_renderer import render_resume as _render_web
        resume = Resume.query.filter_by(slug=slug, is_published=True).first()
        if not resume:
            return render_template("404.html") if _template_exists("404.html") else ("页面不存在", 404)

        # 密码保护检查
        if resume.password_hash:
            token = request.args.get("token", "")
            if not _verify_page_token(slug, token):
                return render_template("web_resume_password.html")

        # 统计浏览量
        resume.view_count = (resume.view_count or 0) + 1
        try:
            from backend.extensions import db as _db
            _db.session.commit()
        except Exception:
            pass

        # 服务端渲染
        html = _render_web(resume.resume_data, resume.web_config)
        return html

    @app.route("/r/<slug>/verify", methods=["POST"])
    def web_resume_verify(slug):
        from backend.models.resume import Resume
        resume = Resume.query.filter_by(slug=slug, is_published=True).first()
        if not resume:
            return jsonify({"success": False, "message": "页面不存在"}), 404
        data = request.get_json() or {}
        password = data.get("password", "")
        if resume.check_password(password):
            token = _generate_page_token(slug)
            return jsonify({"success": True, "data": {"token": token}})
        return jsonify({"success": False, "message": "密码错误"}), 403

    def _template_exists(name):
        import os
        return os.path.exists(os.path.join(app.template_folder or "templates", name))

    def _generate_page_token(slug):
        import hashlib, time as _t
        raw = f"{slug}:{app.config['SECRET_KEY']}:{int(_t.time()) // 3600}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]

    def _verify_page_token(slug, token):
        if not token:
            return False
        import hashlib, time as _t
        for offset in [0, -1]:  # 允许 1 小时误差
            raw = f"{slug}:{app.config['SECRET_KEY']}:{int(_t.time()) // 3600 + offset}"
            expected = hashlib.sha256(raw.encode()).hexdigest()[:16]
            if token == expected:
                return True
        return False

    # ── 旧 API 路由（完整保留，兼容现有前端）──
    @app.route("/api/analyze-resume", methods=["POST"])
    def api_analyze_resume():
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        data = request.get_json()
        if not data or not data.get("resume_text"):
            return jsonify({"error": "请输入简历内容"}), 400
        try:
            return jsonify(analyze_resume(data["resume_text"]))
        except Exception as e:
            return jsonify({"error": "分析失败: %s" % str(e)}), 500

    @app.route("/api/match-jd", methods=["POST"])
    def api_match_jd():
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        data = request.get_json()
        if not data or not data.get("resume_text") or not data.get("jd_text"):
            return jsonify({"error": "请输入简历和职位描述"}), 400
        try:
            return jsonify(match_jd(data["resume_text"], data["jd_text"]))
        except Exception as e:
            return jsonify({"error": "分析失败: %s" % str(e)}), 500

    @app.route("/api/career-advisor", methods=["POST"])
    def api_career_advisor():
        """AI 职业顾问"""
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        data = request.get_json()
        if not data or not data.get("user_query"):
            return jsonify({"error": "请描述你的情况"}), 400
        try:
            if data.get("type") == "offer_compare":
                prompt = '''你是一位资深的职场导师与猎头。用户面临几个Offer的选择。
请根据用户的描述，进行客观的量化性价比分析（考虑时薪、通勤折损、成长性等），指出每个Offer的优劣，并最终给出明确的选择建议。
返回格式严格遵守JSON，不要其他文字：
{
  "analysis": "你的详细分析和建议（包含换行符以分段）。"
}

Offer详情：
''' + data["user_query"]
            else:
                prompt = '''你是一位资深职业规划顾问。根据用户描述，给出职业方向和薪资建议。
你必须且只能输出以下JSON格式，不要输出任何其他文字：
{
  "salary_evaluation": "针对用户期望薪资的合理性评估（50字内）",
  "recommended_roles": [
    {"title": "推荐岗位1", "reason": "推荐理由（30字内）"},
    {"title": "推荐岗位2", "reason": "推荐理由（30字内）"},
    {"title": "推荐岗位3", "reason": "推荐理由（30字内）"}
  ],
  "best_match_jd": "根据用户特征生成的最匹配的一份详细岗位JD，包含岗位职责和任职要求，至少200字"
}

用户描述：
''' + data["user_query"]
            raw = call_gemini(prompt)
            import json as _json
            text = raw.strip()
            if text.startswith("```"):
                text = text.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            result = _json.loads(text)
            return jsonify(result)
        except (ValueError, _json.JSONDecodeError):
            return jsonify({"error": "AI 返回格式异常，请重试"}), 500
        except Exception as e:
            return jsonify({"error": "咨询失败: %s" % str(e)}), 500

    @app.route("/api/inject-keywords", methods=["POST"])
    def api_inject_keywords():
        """AI 智能融合缺失关键词"""
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        data = request.get_json()
        if not data or not data.get("resume_text") or not data.get("missing_keywords"):
            return jsonify({"error": "参数不完整"}), 400
        try:
            keywords = ", ".join(data["missing_keywords"])
            prompt = """你是一位专业的简历优化专家。请将以下缺失的关键词自然地融入简历中。
严格要求：
1. 不得捏造虚假经历或技能，只在合理的位置补充关键词
2. 优先融入自我评价、技能特长和工作职责描述中
3. 保持原文结构和真实信息不变
4. 融入后的表述必须通顺自然，不能生硬堆砌
5. 只输出修改后的完整简历文本，不要加任何解释

需要融入的关键词：""" + keywords + """

原始简历：
""" + data["resume_text"]
            result = call_gemini(prompt)
            return jsonify({"enhanced_resume": result.strip()})
        except Exception as e:
            return jsonify({"error": "融合失败: %s" % str(e)}), 500

    @app.route("/api/parse-resume-pdf", methods=["POST"])
    def api_parse_resume_pdf():
        """解析 PDF 简历并结构化为 JSON"""
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        if "file" not in request.files:
            return jsonify({"error": "请上传文件"}), 400
        f = request.files["file"]
        if not f.filename:
            return jsonify({"error": "文件为空"}), 400
        try:
            import fitz
            pdf_bytes = f.read()
            doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
            text = ""
            for page in doc_pdf:
                text += page.get_text()
            doc_pdf.close()
            if len(text.strip()) < 10:
                return jsonify({"error": "PDF 内容为空或无法识别"}), 400

            prompt = '''你是简历解析专家。请将以下简历纯文本严格解析为JSON格式。
你必须且只能输出JSON，不要输出任何其他文字、解释或markdown标记。

JSON格式要求：
{
  "basic": {"name":"姓名","age":"年龄","phone":"手机","email":"邮箱","city":"城市","years":"工作年限"},
  "intent": {"job":"目标职位","salary":"期望薪资"},
  "education": [{"school":"学校","major":"专业","degree":"学历","time":"时间"}],
  "work": [{"company":"公司","title":"职位","time":"时间","duties":["工作成就1","工作成就2"]}],
  "skills": "技能1、技能2、技能3",
  "intro": "自我评价",
  "certs": "证书"
}

注意：
- 如果某个字段在简历中找不到，填空字符串
- work.duties 必须是字符串数组，每条成就独立一项
- education 和 work 都是数组，支持多段
- skills 用顿号分隔

简历原文：
''' + text[:4000]
            raw = call_gemini(prompt)
            import json as _json
            cleaned = raw.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            result = _json.loads(cleaned)
            if "basic" not in result:
                result = {"basic": {}, "intent": {}, "education": [], "work": [], "skills": "", "intro": "", "certs": ""}
            return jsonify({"resumeData": result})
        except Exception as e:
            return jsonify({"error": "解析失败: %s" % str(e)}), 500

    @app.route("/api/auto-fill", methods=["POST"])
    def api_auto_fill():
        """AI 智能扩写"""
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        data = request.get_json()
        if not data or not data.get("text"):
            return jsonify({"error": "请输入内容"}), 400
        try:
            context = data.get("context", "")
            prompt = """你是一位世界500强企业的HRBP和简历撰写专家。
请将用户的白话文/短句用STAR法则（情境Situation、任务Task、行动Action、结果Result）进行专业扩写。
要求：
1. 主动使用专业动词（主导、统筹、推动、优化、搭建、落地）
2. 补充合理的量化数据（提升X%、降低X%、覆盖X+用户）
3. 每条成就独立一行，不要编号，不要加任何前缀符号
4. 输出3-5条扩写后的成就描述
5. 只输出扩写结果，不要加任何解释文字

""" + ("岗位背景：" + context + "\n\n" if context else "") + "用户原文：\n" + data["text"]
            result = call_gemini(prompt)
            return jsonify({"expanded": result.strip()})
        except Exception as e:
            return jsonify({"error": "扩写失败: %s" % str(e)}), 500

    @app.route("/api/polish-resume", methods=["POST"])
    def api_polish_resume():
        if _is_rate_limited(request.remote_addr):
            return jsonify({"error": "请求过于频繁"}), 429
        data = request.get_json()
        if not data or not data.get("resume_text"):
            return jsonify({"error": "请输入简历内容"}), 400
        try:
            prompt = "你是专业简历润色专家。请润色以下简历，要求：保持信息不变但优化表达，用STAR法则改写工作经历，突出量化成果，技能描述专业化，自我评价精炼有力。只输出润色后的完整简历文本，不要加任何解释。\n\n原始简历：\n" + data["resume_text"]
            result = call_gemini(prompt)
            return jsonify({"polished_text": result})
        except Exception as e:
            return jsonify({"error": "润色失败: %s" % str(e)}), 500

    @app.route("/api/export-word", methods=["POST"])
    def api_export_word():
        data = request.get_json()
        if not data:
            return jsonify({"error": "缺少简历数据"}), 400
        try:
            doc_word = docx.Document()
            style = doc_word.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            basic = data.get('basic', {})
            intent = data.get('intent', {})

            name = basic.get('name', '未命名')
            head = doc_word.add_heading(name, level=0)
            head.alignment = 1

            contact_parts = []
            if basic.get('phone'): contact_parts.append(basic['phone'])
            if basic.get('email'): contact_parts.append(basic['email'])
            if basic.get('city'): contact_parts.append(basic['city'])
            if basic.get('years'): contact_parts.append(f"{basic['years']}经验")
            if intent.get('job'): contact_parts.append(f"求职意向: {intent['job']}")
            if intent.get('salary'): contact_parts.append(f"期望薪资: {intent['salary']}")

            contact_para = doc_word.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = 1

            def add_section(title, content_func):
                doc_word.add_heading(title, level=1)
                content_func()

            if data.get('intro'):
                def add_intro():
                    doc_word.add_paragraph(data['intro'])
                add_section("个人简介", add_intro)

            if data.get('work') and isinstance(data['work'], list):
                def add_work():
                    for w in data['work']:
                        p = doc_word.add_paragraph()
                        p.add_run(w.get('company', '')).bold = True
                        if w.get('title'):
                            p.add_run(f" | {w.get('title')}")
                        if w.get('time'):
                            p.add_run(f"    ({w.get('time')})").italic = True
                        for duty in w.get('duties', []):
                            if str(duty).strip():
                                doc_word.add_paragraph(str(duty).strip(), style='List Bullet')
                add_section("工作经历", add_work)

            if data.get('education') and isinstance(data['education'], list):
                def add_edu():
                    for e in data['education']:
                        p = doc_word.add_paragraph()
                        p.add_run(e.get('school', '')).bold = True
                        parts = []
                        if e.get('major'): parts.append(e['major'])
                        if e.get('degree'): parts.append(e['degree'])
                        if e.get('time'): parts.append(e['time'])
                        if parts:
                            p.add_run(" - " + ", ".join(parts))
                add_section("教育背景", add_edu)

            if data.get('skills'):
                def add_skills():
                    doc_word.add_paragraph(data['skills'])
                add_section("专业技能", add_skills)

            if data.get('certs'):
                def add_certs():
                    doc_word.add_paragraph(data['certs'])
                add_section("证书与荣誉", add_certs)

            file_stream = io.BytesIO()
            doc_word.save(file_stream)
            file_stream.seek(0)

            return send_file(
                file_stream,
                as_attachment=True,
                download_name=f"{name}_简历.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            return jsonify({"error": "Word 生成失败: %s" % str(e)}), 500

    # ── 旧的算力/绑定/简历 API（兼容现有前端）──
    from services.database import (
        get_or_create_user, get_credits, consume,
        add_review, get_public_reviews,
        bind_account, find_user_by_contact, record_login,
        save_resume, list_resumes, load_resume, load_default_resume, delete_resume,
    )
    from services.verify import validate_target, can_send, create_code, verify_code
    from services.database import _conn as _legacy_conn

    @app.route("/api/credits/init", methods=["POST"])
    def api_credits_init():
        data = request.get_json()
        uid = data.get("user_id", "") if data else ""
        device = data.get("device_hash", "") if data else ""
        ip = request.remote_addr or ""
        if not uid:
            return jsonify({"error": "缺少 user_id"}), 400
        user, is_new, err = get_or_create_user(uid, device, ip)
        if err:
            return jsonify({"error": err, "need_bind": True}), 403
        return jsonify({
            "credits": user["credits_left"],
            "is_new": is_new,
            "bound": bool(user.get("bind_email") or user.get("bind_phone")),
        })

    @app.route("/api/credits/consume", methods=["POST"])
    def api_credits_consume():
        data = request.get_json()
        uid = data.get("user_id", "") if data else ""
        if not uid:
            return jsonify({"error": "缺少 user_id"}), 400
        ok, remaining, err = consume(uid)
        if not ok:
            resp = {"error": err, "credits": remaining}
            if "绑定" in (err or ""):
                resp["need_bind"] = True
            return jsonify(resp), 403
        return jsonify({"credits": remaining})

    @app.route("/api/submit-review", methods=["POST"])
    def api_submit_review():
        data = request.get_json()
        if not data:
            return jsonify({"error": "缺少数据"}), 400
        uid = data.get("user_id", "")
        rating = data.get("rating", 0)
        content = data.get("content", "")
        is_anon = 1 if data.get("is_anonymous") else 0
        name = data.get("display_name", "匿名用户")
        feat = data.get("feature", "general")
        if not uid or not rating:
            return jsonify({"error": "缺少评分"}), 400
        new_credits, err = add_review(uid, feat, rating, content, is_anon, name)
        if err:
            return jsonify({"error": err}), 400
        bound = bool(find_user_by_contact(email=None, phone=None) is not None)
        return jsonify({
            "credits": new_credits,
            "message": "感谢评价！⚡+100 算力",
            "show_bind": not bound,
        })

    @app.route("/api/reviews", methods=["GET"])
    def api_get_reviews():
        limit = request.args.get("limit", 20, type=int)
        return jsonify({"reviews": get_public_reviews(limit)})

    @app.route("/api/verify/send", methods=["POST"])
    def api_verify_send():
        data = request.get_json()
        target = data.get("target", "") if data else ""
        purpose = data.get("purpose", "bind") if data else "bind"
        t_type, cleaned, err = validate_target(target)
        if err:
            return jsonify({"error": err}), 400
        conn = _legacy_conn()
        ok, err2, wait = can_send(conn, cleaned, purpose)
        if not ok:
            return jsonify({"error": err2, "wait": wait}), 429
        code = create_code(conn, cleaned, purpose)
        if t_type == "email":
            from services.email_sender import send_verify_code as _send
            ok2, err3 = _send(cleaned, code, purpose)
            if not ok2:
                return jsonify({"error": err3}), 500
        else:
            print("[短信] 验证码: %s → %s" % (code, cleaned))
        mask = cleaned[:3] + "****" + cleaned[-4:]
        return jsonify({"message": "验证码已发送至 " + mask})

    @app.route("/api/bind-account", methods=["POST"])
    def api_bind_account():
        data = request.get_json()
        uid = data.get("user_id", "") if data else ""
        email = data.get("email")
        phone = data.get("phone")
        code = data.get("code", "")
        if not uid:
            return jsonify({"error": "缺少 user_id"}), 400
        if not email and not phone:
            return jsonify({"error": "请输入邮箱或手机号"}), 400
        target = email or phone
        conn = _legacy_conn()
        valid, err = verify_code(conn, target, code, "bind")
        if not valid:
            return jsonify({"error": err}), 400
        result, err2 = bind_account(uid, email=email, phone=phone)
        if err2:
            return jsonify({"error": err2}), 400
        return jsonify(result)

    @app.route("/api/auth/login", methods=["POST"])
    def api_legacy_login():
        data = request.get_json()
        target = data.get("target", "") if data else ""
        code = data.get("code", "") if data else ""
        if not target or not code:
            return jsonify({"error": "请输入账号和验证码"}), 400
        t_type, cleaned, err = validate_target(target)
        if err:
            return jsonify({"error": err}), 400
        conn = _legacy_conn()
        valid, err2 = verify_code(conn, cleaned, code, "login")
        if not valid:
            return jsonify({"error": err2}), 400
        user = find_user_by_contact(
            email=cleaned if t_type == "email" else None,
            phone=cleaned if t_type == "phone" else None
        )
        if not user:
            return jsonify({"error": "该账号尚未绑定，请先注册或绑定"}), 404
        record_login(user["user_id"])
        default_resume = load_default_resume(user["user_id"])
        return jsonify({
            "user_id": user["user_id"],
            "credits": user["credits_left"],
            "bound": True,
            "default_resume": default_resume,
        })

    @app.route("/api/resumes/save", methods=["POST"])
    def api_resumes_save():
        data = request.get_json()
        uid = data.get("user_id", "") if data else ""
        if not uid:
            return jsonify({"error": "缺少 user_id", "need_bind": True}), 400
        import json as _j
        rid, err = save_resume(
            uid,
            _j.dumps(data.get("resume_data", {}), ensure_ascii=False),
            _j.dumps(data.get("template_config", {}), ensure_ascii=False),
            title=data.get("title", "未命名简历"),
            resume_id=data.get("resume_id"),
            is_default=data.get("is_default", False),
        )
        if err:
            return jsonify({"error": err}), 400
        return jsonify({"resume_id": rid})

    @app.route("/api/resumes/list", methods=["GET"])
    def api_resumes_list():
        uid = request.args.get("user_id", "")
        if not uid:
            return jsonify({"error": "缺少 user_id"}), 400
        return jsonify({"resumes": list_resumes(uid)})

    @app.route("/api/resumes/load", methods=["GET"])
    def api_resumes_load():
        uid = request.args.get("user_id", "")
        rid = request.args.get("resume_id", type=int)
        if not uid or not rid:
            return jsonify({"error": "参数不完整"}), 400
        data, err = load_resume(uid, rid)
        if err:
            return jsonify({"error": err}), 404
        return jsonify(data)

    @app.route("/api/resumes/delete", methods=["POST"])
    def api_resumes_delete():
        data = request.get_json()
        uid = data.get("user_id", "") if data else ""
        rid = data.get("resume_id") if data else None
        if not uid or not rid:
            return jsonify({"error": "参数不完整"}), 400
        ok, err = delete_resume(uid, rid)
        if not ok:
            return jsonify({"error": err}), 400
        return jsonify({"success": True})

    return app


def _setup_scheduler(app):
    """配置定时任务"""
    try:
        from apscheduler.schedulers.background import BackgroundScheduler

        scheduler = BackgroundScheduler()

        def check_expired_subscriptions():
            with app.app_context():
                from backend.extensions import db
                from backend.models.subscription import Subscription
                from datetime import datetime, timezone

                now = datetime.now(timezone.utc)
                expired = Subscription.query.filter(
                    Subscription.status == "active",
                    Subscription.current_period_end < now,
                ).all()

                for sub in expired:
                    if sub.auto_renew:
                        sub.status = "past_due"
                    else:
                        sub.status = "expired"
                    sub.updated_at = now

                if expired:
                    db.session.commit()
                    print(f"[定时] 已处理 {len(expired)} 个到期订阅")

        def cleanup_verify_codes():
            with app.app_context():
                from backend.extensions import db
                from backend.models.verify_code import VerifyCode
                from datetime import datetime, timezone, timedelta

                cutoff = datetime.now(timezone.utc) - timedelta(hours=24)
                deleted = VerifyCode.query.filter(VerifyCode.created_at < cutoff).delete()
                if deleted:
                    db.session.commit()
                    print(f"[定时] 清理 {deleted} 个过期验证码")

        scheduler.add_job(check_expired_subscriptions, 'interval', hours=1, id='check_subscriptions')
        scheduler.add_job(cleanup_verify_codes, 'interval', hours=6, id='cleanup_codes')

        scheduler.start()
        print("[ok] 定时任务已启动")
    except Exception as e:
        print(f"[提示] 定时任务启动失败: {e}")


# ── 入口 ──
app = create_app()

if __name__ == "__main__":
    print("服务启动于: http://localhost:%d" % PORT)
    app.run(host="0.0.0.0", port=PORT, debug=DEBUG)